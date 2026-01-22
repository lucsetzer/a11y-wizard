# pdf_analyzer.py
import os
import tempfile
from typing import Dict, List, Any
import PyPDF2
from docx import Document
import re

class PDFAccessibilityChecker:
    """Check PDF and Word documents for accessibility issues"""
    
    def extract_document_text(self, file_path: str, filename: str) -> str:
        """Extract text content from document"""
        file_ext = os.path.splitext(filename)[1].lower()
        
        try:
            if file_ext == '.pdf':
                text = []
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for i, page in enumerate(pdf_reader.pages[:5]):  # First 5 pages only
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text.append(f"--- Page {i+1} ---\n{page_text[:1000]}")  # Limit per page
                return "\n\n".join(text) if text else "No extractable text found"
            
            elif file_ext in ['.doc', '.docx']:
                doc = Document(file_path)
                paragraphs = []
                for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
                    if para.text and para.text.strip():
                        paragraphs.append(para.text[:500])  # Limit per paragraph
                return "\n".join(paragraphs) if paragraphs else "No text content found"
            
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read(2000)  # First 2000 chars
            
            return "Unsupported file type for text extraction"
        
        except Exception as e:
            return f"Error extracting text: {str(e)[:200]}"
        
    
    def analyze_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Analyze a document for accessibility issues"""
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == '.pdf':
            return self._analyze_pdf(file_path, filename)
        elif file_ext in ['.doc', '.docx']:
            return self._analyze_word(file_path, filename)
        elif file_ext == '.txt':
            return self._analyze_text(file_path, filename)
        else:
            return self._error_result(filename, f"Unsupported file type: {file_ext}")
    
        def extract_document_text(self, file_path: str, filename: str) -> str:
            """Extract text content from document for AI analysis"""
            file_ext = os.path.splitext(filename)[1].lower()
            
            try:
                if file_ext == '.pdf':
                    text = []
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for i, page in enumerate(pdf_reader.pages[:5]):  # First 5 pages only
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                text.append(f"--- Page {i+1} ---\n{page_text[:1000]}")  # Limit per page
                    return "\n\n".join(text) if text else "No extractable text found"
                
                elif file_ext in ['.doc', '.docx']:
                    doc = Document(file_path)
                    paragraphs = []
                    for i, para in enumerate(doc.paragraphs[:50]):  # First 50 paragraphs
                        if para.text and para.text.strip():
                            paragraphs.append(para.text[:500])  # Limit per paragraph
                    return "\n".join(paragraphs) if paragraphs else "No text content found"
                
                elif file_ext == '.txt':
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                        return file.read(2000)  # First 2000 chars
                
                return "Unsupported file type for text extraction"
            
            except Exception as e:
                return f"Error extracting text: {str(e)[:200]}"
                
        
        def _analyze_word(self, file_path: str, filename: str) -> Dict[str, Any]:
            """Analyze Word document for accessibility issues - ENHANCED"""
            issues = []
        
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            # 1. Check for document title
            if not core_props.title:
                issues.append({
                    "type": "critical",
                    "title": "Missing Document Title",
                    "count": 1,
                    "description": "Word document missing title property",
                    "fix": "Add title in File → Properties → Title field",
                    "category": "Document Structure",
                    "example": "Title: 'Research Paper - Climate Change Impact'"
                })
            
            # 2. Check heading structure
            headings = []
            heading_levels = set()
            
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                if style_name.startswith('Heading'):
                    headings.append(style_name)
                    # Extract heading level
                    try:
                        level = int(style_name.split()[-1])
                        heading_levels.add(level)
                    except:
                        pass
            
            if not headings:
                issues.append({
                    "type": "warning",
                    "title": "No Heading Structure",
                    "count": 1,
                    "description": "Document doesn't use heading styles",
                    "fix": "Use Heading 1, Heading 2, etc. styles for document structure",
                    "category": "Structure",
                    "example": "Apply styles from Home → Styles gallery"
                })
            else:
                # Check if Heading 1 exists
                if 1 not in heading_levels:
                    issues.append({
                        "type": "warning",
                        "title": "Missing H1 Heading",
                        "count": 1,
                        "description": "No Heading 1 style found (main document title)",
                        "fix": "Use Heading 1 for main document title",
                        "category": "Structure"
                    })
                
                # Check for skipped heading levels
                if heading_levels and max(heading_levels) > 1:
                    for i in range(1, max(heading_levels)):
                        if i not in heading_levels and i+1 in heading_levels:
                            issues.append({
                                "type": "warning",
                                "title": "Skipped Heading Level",
                                "count": 1,
                                "description": f"Heading level {i} skipped (went from H{i-1} to H{i+1})",
                                "fix": "Maintain sequential heading levels",
                                "category": "Structure"
                            })
                            break
            
            # 3. Check for alt text on images
            images_without_alt = []
            for shape in doc.inline_shapes:
                # Check if shape has alt text (basic check)
                # Note: python-docx doesn't have direct alt text access
                images_without_alt.append(shape)
            
            image_count = len(doc.inline_shapes)
            if image_count > 0:
                issues.append({
                    "type": "warning" if image_count < 5 else "critical",
                    "title": "Images Need Alt Text Check",
                    "count": image_count,
                    "description": f"Document contains {image_count} image(s) - verify alt text",
                    "fix": "Right-click each image → Edit Alt Text → add description",
                    "category": "Images",
                    "example": "Alt text: 'Bar chart showing quarterly sales growth'"
                })
            
            # 4. Check for tables
            table_count = len(doc.tables)
            if table_count > 0:
                tables_without_headers = []
                for table in doc.tables:
                    # Basic check: first row should be different (header)
                    if len(table.rows) > 0:
                        first_row = table.rows[0]
                        # Simple heuristic: check if first cell has bold or different style
                        has_header_style = False
                        if first_row.cells:
                            cell_text = first_row.cells[0].text.strip()
                            if cell_text and cell_text.upper() == cell_text:  # ALL CAPS often indicates header
                                has_header_style = True
                        
                        if not has_header_style:
                            tables_without_headers.append(table)
                
                if tables_without_headers:
                    issues.append({
                        "type": "warning",
                        "title": "Tables May Need Headers",
                        "count": len(tables_without_headers),
                        "description": f"{len(tables_without_headers)} table(s) may need header rows",
                        "fix": "Designate first row as header row (Table Design → Header Row)",
                        "category": "Tables"
                    })
            
            # 5. Check for hyperlinks
            hyperlink_count = 0
            for paragraph in doc.paragraphs:
                if 'HYPERLINK' in paragraph._element.xml:
                    hyperlink_count += 1
            
            if hyperlink_count > 0:
                issues.append({
                    "type": "info",
                    "title": "Contains Hyperlinks",
                    "count": hyperlink_count,
                    "description": f"Document has {hyperlink_count} hyperlink(s)",
                    "fix": "Ensure link text is descriptive (not 'click here')",
                    "category": "Links"
                })
            
            # 6. Check document properties
            if not core_props.author:
                issues.append({
                    "type": "info",
                    "title": "Missing Author",
                    "count": 1,
                    "description": "Document missing author information",
                    "fix": "Add author in File → Properties",
                    "category": "Metadata"
                })
            
            # Calculate score
            score = self._calculate_score(issues, len(doc.paragraphs))
            
            # Generate summary
            summary_parts = []
            if headings:
                summary_parts.append(f"{len(headings)} heading(s)")
            if image_count:
                summary_parts.append(f"{image_count} image(s)")
            if table_count:
                summary_parts.append(f"{table_count} table(s)")
            
            summary = f"Word document: {len(issues)} issues found"
            if summary_parts:
                summary += f" with {', '.join(summary_parts)}"
            
            return {
                "filename": filename,
                "score": score,
                "issues": issues,
                "document_info": {
                    "paragraphs": len(doc.paragraphs),
                    "headings": len(headings),
                    "heading_levels": list(heading_levels),
                    "images": image_count,
                    "tables": table_count,
                    "hyperlinks": hyperlink_count,
                    "author": core_props.author or "Missing",
                    "title": core_props.title or "Missing",
                    "pages": getattr(core_props, 'pages', 'Unknown')
                },
                "summary": summary,
                "method": "word-analysis-enhanced"
            }
            
        except Exception as e:
            return self._error_result(filename, f"Word analysis error: {str(e)}")
    
    def _analyze_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Analyze plain text file"""
        issues = []
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                lines = content.split('\n')
                
                # Very basic text file checks
                if len(content) > 1000 and '\n\n\n' in content:
                    issues.append({
                        "type": "warning",
                        "title": "Poor Paragraph Spacing",
                        "count": 1,
                        "description": "Text file has excessive blank lines",
                        "fix": "Use consistent single blank lines between paragraphs",
                        "category": "Formatting"
                    })
                
                # Check line length
                long_lines = [line for line in lines if len(line) > 100]
                if long_lines:
                    issues.append({
                        "type": "warning",
                        "title": "Long Lines",
                        "count": len(long_lines),
                        "description": f"{len(long_lines)} line(s) exceed 100 characters",
                        "fix": "Break long lines for better readability",
                        "category": "Readability"
                    })
                
                score = self._calculate_score(issues, len(lines))
                
                return {
                    "filename": filename,
                    "score": score,
                    "issues": issues,
                    "line_count": len(lines),
                    "char_count": len(content),
                    "summary": f"Text file analysis: {len(issues)} issues found",
                    "method": "text-analysis"
                }
                
        except Exception as e:
            return self._error_result(filename, f"Text analysis error: {str(e)}")
    
       
    def _calculate_score(self, issues, element_count):
        """Calculate accessibility score for documents"""
        # Start with 100
        score = 100
        
        # Deduct for issues
        for issue in issues:
            if issue.get("type") == "critical":
                score -= 15
            elif issue.get("type") == "warning":
                score -= 5
        
        # Bonus for having content
        if element_count > 10:
            score += 5
        
        # Ensure score is reasonable
        if score < 30:
            score = 30
        if score > 100:
            score = 100
            
        return score
    
    def _error_result(self, filename: str, error: str):
        """Create error result"""
        return {
            "filename": filename,
            "score": 0,
            "issues": [{
                "type": "critical",
                "title": "Analysis Failed",
                "count": 1,
                "description": error[:100],
                "fix": "Try a different file or format",
                "category": "Error"
            }],
            "summary": f"Analysis failed: {error[:50]}",
            "method": "error"
        }


    def _error_result(self, filename: str, error: str) -> Dict[str, Any]:
        """Create error result"""
        return {
            "filename": filename,
            "score": 0,
            "issues": [{
                "type": "critical",
                "title": "Analysis Failed",
                "count": 1,
                "description": error[:100],
                "fix": "Try a different file or format",
                "category": "Error"
            }],
            "summary": f"Analysis failed: {error[:50]}",
            "method": "error"
        }
    
    def _analyze_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Analyze PDF for accessibility issues - SIMPLE VERSION"""
        issues = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Basic checks
                if not pdf_reader.metadata or not pdf_reader.metadata.get('/Title'):
                    issues.append({
                        "type": "critical",
                        "title": "Missing Document Title",
                        "count": 1,
                        "description": "PDF missing title in document properties",
                        "fix": "Add a title in PDF properties (File → Properties)",
                        "category": "Document Structure"
                    })
                
                # Check for bookmarks/outline (navigation)
                if not pdf_reader.outline:
                    issues.append({
                        "type": "warning",
                        "title": "No Document Bookmarks",
                        "count": 1,
                        "description": "PDF lacks bookmarks for navigation",
                        "fix": "Add bookmarks for major sections",
                        "category": "Navigation"
                    })
                
                # Check page count
                page_count = len(pdf_reader.pages)
                
                # Try to extract text
                has_text = False
                for page in pdf_reader.pages[:3]:  # Check first 3 pages
                    text = page.extract_text()
                    if text and len(text.strip()) > 0:
                        has_text = True
                        break
                
                if not has_text:
                    issues.append({
                        "type": "critical",
                        "title": "Scanned/Image PDF",
                        "count": 1,
                        "description": "PDF appears to be scanned images without selectable text",
                        "fix": "Use OCR to create searchable text",
                        "category": "Text"
                    })
                
                # Calculate score
                score = self._calculate_score(issues, page_count)
                
                return {
                    "filename": filename,
                    "score": score,
                    "issues": issues,
                    "page_count": page_count,
                    "has_text": has_text,
                    "summary": f"PDF analysis: {len(issues)} issues found across {page_count} pages",
                    "method": "pdf-analysis"
                }
                
        except Exception as e:
            return self._error_result(filename, f"PDF analysis error: {str(e)}")
